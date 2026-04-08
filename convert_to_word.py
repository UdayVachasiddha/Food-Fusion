import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []
    in_mermaid = False

    for line in lines:
        clean_line = line.strip()

        # Handle Mermaid Blocks
        if clean_line.startswith('```mermaid'):
            in_mermaid = True
            p = doc.add_paragraph("--- DATABASE ERD (MERMAID CODE) ---")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if clean_line.startswith('```') and in_mermaid:
            in_mermaid = False
            continue
        if in_mermaid:
            p = doc.add_paragraph(clean_line)
            p.style = 'No Spacing'
            font = p.runs[0].font if p.runs else None
            if font: font.name = 'Courier New'
            continue

        # Handle Horizontal Rules
        if clean_line == '---':
            doc.add_page_break()
            continue

        # Handle Tables
        if '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table:
            # Construct the table in Word
            rows = []
            for tl in table_lines:
                if '---' in tl: continue # Skip separator line
                cells = [c.strip() for c in tl.split('|') if c.strip() or tl.startswith('|') or tl.endswith('|')]
                # Filter out empty cells caused by leading/trailing pipes
                cells = [c for c in cells if c != '']
                if len(cells) > 0:
                    rows.append(cells)
            
            if rows:
                word_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                word_table.style = 'Table Grid'
                for r_idx, r_data in enumerate(rows):
                    for c_idx, c_text in enumerate(r_data):
                        cell = word_table.cell(r_idx, c_idx)
                        # Check for images inside cell
                        img_match = re.search(r'!\[.*?\]\((.*?)\)', c_text)
                        if img_match:
                            img_path = img_match.group(1).replace('file:///', '').replace('/', '\\')
                            if os.path.exists(img_path):
                                run = cell.paragraphs[0].add_run()
                                try:
                                    run.add_picture(img_path, width=Inches(3.0))
                                except Exception as e:
                                    cell.paragraphs[0].text = f"[Image Error: {os.path.basename(img_path)}]"
                            else:
                                cell.paragraphs[0].text = f"[Image Not Found: {os.path.basename(img_path)}]"
                        else:
                            # Handle clean text (removing bold markers for cleaner Word look)
                            clean_text = c_text.replace('**', '')
                            cell.text = clean_text
            
            in_table = False
            table_lines = []
            continue

        # Handle Headings
        if clean_line.startswith('# '):
            doc.add_heading(clean_line[2:], level=1)
        elif clean_line.startswith('## '):
            doc.add_heading(clean_line[3:], level=2)
        elif clean_line.startswith('### '):
            doc.add_heading(clean_line[4:], level=3)
        
        # Handle Paragraphs and Bold text
        elif clean_line:
            p = doc.add_paragraph()
            # Basic bold parsing: **text**
            parts = re.split(r'(\*\*.*?\*\*)', clean_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    import traceback
    print("Main block entered")
    try:
        md_file = r"d:\PHP\htdocs\BE Assignment\Task5_Reflection.md"
        print(f"MD File: {md_file}")
        word_file = r"d:\PHP\htdocs\BE Assignment\Task5_Reflection.docx"
        convert_md_to_docx(md_file, word_file)
    except Exception:
        traceback.print_exc()
